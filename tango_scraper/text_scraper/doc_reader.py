from docx import Document
import os

from data_format import Page


class DocReader:

    def get_text_from_docx(self, file):
        """Extract text from a DOCX file.
            args: file
            returns: pages
        """
        print("Extracting text from DOCX file")
        document = Document(file)
        pages: list[Page] = []
        for i, para in enumerate(document.paragraphs):
            pages.append(
                Page(text=para.text, page_num=i, pdf_name=os.path.basename(file.name))
            )
        return pages

    def get_text_from_doc(self, file):
        """Extract text from a DOC file, assuming it has been converted to DOCX format."""
        print("Extracting text from DOC file")
        if not file.name.endswith(".docx"):
            return "This file is not a valid DOC file. Please convert to DOCX format before processing."
        return self.get_text_from_docx(file)

    def extract_text_from_docx(self, file):
        """Extract text from a DOCX file."""
        if file.name.endswith(".docx"):
            document = Document(file)
            pages: list[Page] = []
            for i, para in enumerate(document.paragraphs):
                pages.append(
                    Page(
                        text=para.text, page_num=i, pdf_name=os.path.basename(file.name)
                    )
                )

            return pages
        else:
            return "This file is not a valid DOCX file."

    def get_text_from_doc_or_docx(self, file):
        """Determine the file type and extract text accordingly."""
        print("Extracting text from DOC or DOCX file")
        _, file_extension = os.path.splitext(file.name)
        if file_extension.lower() == ".docx":
            return self.get_text_from_docx(file)
        elif file_extension.lower() == ".doc":
            return self.get_text_from_doc(file)
        else:
            raise ValueError("Unsupported file type")
